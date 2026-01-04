from langchain.schema import Document
from pypdf import PdfReader
import pandas as pd
import json
import docx
import os

def load_documents(path: str):
    docs = []
    ext = os.path.splitext(path)[1]

    if ext == ".pdf":
        reader = PdfReader(path)
        for page in reader.pages:
            docs.append(Document(page_content=page.extract_text()))
    
    elif ext == ".csv":
        df = pd.read_csv(path)
        for _, row in df.iterrows():
            docs.append(Document(page_content=row.to_string()))

    elif ext == ".json":
        with open(path) as f:
            data = json.load(f)
        docs.append(Document(page_content=json.dumps(data)))

    elif ext == ".docx":
        doc = docx.Document(path)
        docs.append(Document(page_content="\n".join(p.text for p in doc.paragraphs)))

    elif ext == ".txt":
        with open(path) as f:
            docs.append(Document(page_content=f.read()))

    return docs
